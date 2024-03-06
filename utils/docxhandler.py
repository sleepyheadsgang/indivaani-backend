from docx import Document
import os

input_file_path = "/Users/shivam/Desktop/Sleepyheads/input.docx"

output_file_path = "/Users/shivam/Desktop/Sleepyheads/indivaani-backend/out/output.docx"

output_document = Document()

input_document = Document(input_file_path)

for paragraph in input_document.paragraphs:
    output_document.add_paragraph(paragraph.text)

output_document.save(output_file_path)

print("Output document generated successfully at:", output_file_path)
